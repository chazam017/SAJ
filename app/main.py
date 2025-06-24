from fastapi import FastAPI, UploadFile, File, Form, Body
from fastapi.responses import FileResponse, JSONResponse, PlainTextResponse
from fastapi.middleware.cors import CORSMiddleware
from docx import Document
from datetime import datetime
from PIL import Image
from pdf2image import convert_from_bytes
import pytesseract
import shutil
import os
import io
import re

# OCR config
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
POPPLER_PATH = r"C:\poppler\Library\bin"

# inicialização
app = FastAPI()
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# pastas/rotas
SALVOS_DIR = "saida"
UPLOAD_DIR = "uploads"
TEMPLATES_DIR = "templates"
os.makedirs(SALVOS_DIR, exist_ok=True)
os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(TEMPLATES_DIR, exist_ok=True)

# substituição das TAGs
def substituir_tags(doc: Document, dados: dict):
    for p in doc.paragraphs:
        for tag, valor in dados.items():
            p.text = p.text.replace(f"{{{{{tag}}}}}", valor)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for tag, valor in dados.items():
                    cell.text = cell.text.replace(f"{{{{{tag}}}}}", valor)

# gerar a petição com as tags
@app.post("/gerar-peticao")
async def gerar_peticao(
    nome: str = Form(...),
    nascimento: str = Form(...),
    cpf: str = Form(...),
    rg: str = Form(...),
    beneficio: str = Form(...),
    enderecamento: str = Form(...),
    nacionalidade: str = Form(...),
    civil: str = Form(...),
    profissao: str = Form(...),
    filiacao: str = Form(...),
    endereco: str = Form(...),
    procurador: str = Form(...),
    oab: str = Form(...),
    der: str = Form(...),
    nb: str = Form(...),
    documentos: list[UploadFile] = File([]),
):
    try:
        template_map = {
            "idoso": "1. TEMPLATE - LOAS IDOSO NOVO.docx",
            "deficiente": "1. TEMPLATE - LOAS DEFICIENTE NOVO.docx",
            "rural_salario": "TEMPLATE - SALÁRIO MATERNIDADE RURAL.docx",
            "urbano_salario": "TEMPLATE - SALÁRIO MATERNIDADE URBANO.docx"
        }
        template_file = template_map.get(beneficio, "MODELO_PADRAO.docx")
        template_path = os.path.join(TEMPLATES_DIR, template_file)

        doc = Document(template_path)

        dados = {
            "NOME": nome,
            "NASCIMENTO": nascimento,
            "CPF": cpf,
            "RG": rg,
            "ENDERECAMENTO": enderecamento,
            "NACIONALIDADE": nacionalidade,
            "CIVIL": civil,
            "PROFISSAO": profissao,
            "FILIAÇÃO": filiacao,
            "ENDERECO": endereco,
            "PROCURADOR": procurador,
            "OAB": oab,
            "DER": der,
            "NB": nb
        }

        substituir_tags(doc, dados)

        ts = datetime.now().strftime("%Y%m%d%H%M%S")
        filename = f"peticao_{nome.strip().lower().replace(' ', '_')}_{ts}.docx"
        path = os.path.join(SALVOS_DIR, filename)
        doc.save(path)

        return JSONResponse({"success": True, "filename": filename, "url": f"/download/{filename}"})
    except Exception as e:
        return JSONResponse({"error": f"Erro ao gerar petição: {e}"}, status_code=500)

# download da petição finalizada
@app.get("/download/{filename}")
async def download_peticao(filename: str):
    path = os.path.join(SALVOS_DIR, filename)
    if not os.path.exists(path):
        return JSONResponse({"error": "Arquivo não encontrado."}, status_code=404)
    return FileResponse(path, filename=filename, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# pré-visualização da última petição
@app.get("/download/ultimo")
async def download_ultimo_texto():
    try:
        arquivos = sorted(
            [f for f in os.listdir(SALVOS_DIR) if f.endswith(".docx")],
            key=lambda x: os.path.getmtime(os.path.join(SALVOS_DIR, x)),
            reverse=True
        )
        if not arquivos:
            return JSONResponse({"error": "Nenhuma petição encontrada."}, status_code=404)
        doc = Document(os.path.join(SALVOS_DIR, arquivos[0]))
        texto = "\n".join(p.text for p in doc.paragraphs)
        return PlainTextResponse(texto)
    except Exception as e:
        return JSONResponse({"error": f"Erro ao ler petição: {e}"}, status_code=500)

# download do ultimo arquivo
@app.get("/download/ultimo-arquivo")
async def download_ultimo_arquivo():
    arquivos = sorted(
        [f for f in os.listdir(SALVOS_DIR) if f.endswith(".docx")],
        key=lambda x: os.path.getmtime(os.path.join(SALVOS_DIR, x)),
        reverse=True
    )
    if not arquivos:
        return JSONResponse({"error": "Nenhum documento encontrado."}, status_code=404)
    path = os.path.join(SALVOS_DIR, arquivos[0])
    return FileResponse(path, filename=arquivos[0], media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# upload de documentos p/ montar a petição
@app.post("/upload-documentos")
async def upload_documentos(documentos: list[UploadFile] = File(...)):
    try:
        salvos = []
        for f in documentos:
            dest = os.path.join(UPLOAD_DIR, f.filename)
            with open(dest, "wb") as out:
                shutil.copyfileobj(f.file, out)
            salvos.append(f.filename)
        return JSONResponse({"success": True, "arquivos_salvos": salvos})
    except Exception as e:
        return JSONResponse({"error": f"Erro ao salvar arquivos: {e}"}, status_code=500)

# OCR para extração da RMI (cálculo do valor da causa)
@app.get("/extrair-rmi")
async def extrair_rmi():
    arquivos = sorted(
        [f for f in os.listdir(UPLOAD_DIR) if f.lower().endswith((".pdf", ".png", ".jpg", ".jpeg"))],
        key=lambda x: os.path.getmtime(os.path.join(UPLOAD_DIR, x)),
        reverse=True
    )
    if not arquivos:
        return JSONResponse({"rmi": None})

    texto_total = ""
    for nome in arquivos:
        path = os.path.join(UPLOAD_DIR, nome)
        with open(path, "rb") as f:
            content = f.read()
        try:
            if nome.lower().endswith((".png", ".jpg", ".jpeg")):
                img = Image.open(io.BytesIO(content))
                texto_total += pytesseract.image_to_string(img)
            else:
                imagens = convert_from_bytes(content, poppler_path=POPPLER_PATH)
                for img in imagens:
                    texto_total += pytesseract.image_to_string(img)
        except:
            continue

    match = re.search(r"RMI\s*[:\-]?\s*R?\$?\s*(\d+[.,]?\d{0,2})", texto_total)
    if match:
        return JSONResponse({"rmi": match.group(1).replace(",", ".")})
    return JSONResponse({"rmi": None})

# salvar edição manual
@app.post("/salvar-edicao")
async def salvar_edicao(payload: dict = Body(...)):
    txt = payload.get("texto", "").strip()
    if not txt:
        return JSONResponse({"error": "Texto vazio."}, status_code=400)

    ts = datetime.now().strftime("%Y%m%d%H%M%S")
    fn = f"peticao_editada_{ts}.docx"
    path = os.path.join(SALVOS_DIR, fn)

    doc = Document()
    for linha in txt.splitlines():
        doc.add_paragraph(linha)

    doc.save(path)
    return JSONResponse({"success": True, "filename": fn, "url": f"/download/{fn}"})

# listar os modelos já cadastrados
@app.get("/templates")
async def listar_templates():
    files = [f for f in os.listdir(TEMPLATES_DIR) if f.lower().endswith(".docx")]
    return JSONResponse({"templates": files})

# fazer upload de novos modelos de petições
@app.post("/templates/upload")
async def upload_template(template: UploadFile = File(...)):
    if not template.filename.lower().endswith(".docx"):
        return JSONResponse({"error": "Apenas arquivos .docx são permitidos."}, status_code=400)
    dest = os.path.join(TEMPLATES_DIR, template.filename)
    with open(dest, "wb") as out:
        shutil.copyfileobj(template.file, out)
    return JSONResponse({"success": True, "filename": template.filename})

# excluir os modelos
@app.delete("/templates/{filename}")
async def deletar_template(filename: str):
    path = os.path.join(TEMPLATES_DIR, filename)
    if os.path.exists(path):
        os.remove(path)
        return JSONResponse({"success": True})
    return JSONResponse({"error": "Template não encontrado."}, status_code=404)

# editar os modelos dos documentos
@app.post("/templates/editar")
async def editar_template(
    antigo_nome: str = Form(...),
    novo_nome: str = Form(None),
    novo_arquivo: UploadFile = File(None)
):
    try:
        antigo_path = os.path.join(TEMPLATES_DIR, antigo_nome)
        if not os.path.exists(antigo_path):
            return JSONResponse({"error": "Arquivo original não encontrado."}, status_code=404)

        if novo_nome:
            novo_nome = novo_nome if novo_nome.endswith(".docx") else novo_nome + ".docx"
            novo_path = os.path.join(TEMPLATES_DIR, novo_nome)
            os.rename(antigo_path, novo_path)
            return JSONResponse({"success": True, "mensagem": "Template renomeado com sucesso."})

        if novo_arquivo:
            with open(antigo_path, "wb") as out:
                shutil.copyfileobj(novo_arquivo.file, out)
            return JSONResponse({"success": True, "mensagem": "Arquivo substituído com sucesso."})

        return JSONResponse({"error": "Nenhuma alteração recebida."}, status_code=400)
    except Exception as e:
        return JSONResponse({"error": f"Erro ao editar template: {e}"}, status_code=500)